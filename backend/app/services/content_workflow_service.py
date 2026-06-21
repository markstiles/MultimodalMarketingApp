import io
import logging
import os
from datetime import datetime, timezone

import httpx
from docx import Document

from app.services.sitecore_auth import get_sitecore_automation_token

logger = logging.getLogger(__name__)

# ── Phase artifact registry ──────────────────────────────────────────────────

PHASE_ARTIFACT_MAP: dict[str, dict[str, str]] = {
    "Research":   {"filename": "research-brief.docx"},
    "Strategy":   {"filename": "marketing-strategy.docx"},
    "BrandVoice": {"filename": "brand-voice-summary.docx"},
    "Brief":      {"filename": "campaign-brief.docx"},
    "Campaign":   {"filename": "campaign-plan.docx"},
}

CONTENT_STRATEGY_FOLDER = "Content Strategy"

PHASES_ORDERED: list[str] = list(PHASE_ARTIFACT_MAP.keys())
STALENESS_DAYS = 365

# Sitecore template ID for media library folders
FOLDER_TEMPLATE_ID = "FE5DD826-48C6-436D-B87A-7C4210C7413B"


def build_artifact_media_path(tenant: str, site: str, phase: str) -> str:
    info = PHASE_ARTIFACT_MAP.get(phase)
    if not info:
        raise ValueError(f"Unknown phase: {phase!r}")
    # Strip the extension: this instance's InvalidItemNameChars includes '.' so
    # Sitecore creates items without the extension in their path name. The
    # physical filename (with .docx) is used only for the multipart upload.
    filename = info["filename"]
    item_name = filename.rsplit(".", 1)[0] if "." in filename else filename
    return (
        f"/sitecore/Media Library/Project/{tenant}/{site}"
        f"/{CONTENT_STRATEGY_FOLDER}/{item_name}"
    )


# ── Auth ─────────────────────────────────────────────────────────────────────

# Alias for backward compatibility — callers in this module use the old name.
get_sitecore_media_auth_token = get_sitecore_automation_token


# ── Media library existence check ────────────────────────────────────────────

async def check_media_artifact_exists(
    tenant: str, site: str, phase: str, auth_token: str
) -> dict:
    media_path = build_artifact_media_path(tenant, site, phase)
    cm_host = os.environ.get("SITECORE_CM_HOST", "").rstrip("/")
    if not cm_host:
        logger.warning("SITECORE_CM_HOST not set — treating all phases as not_started")
        return {"exists": False, "modified_at": None, "age_days": None}

    url = f"{cm_host}/sitecore/api/ssc/item"
    params = {"database": "master", "path": media_path}

    try:
        async with httpx.AsyncClient(timeout=8) as http:
            resp = await http.get(
                url,
                params=params,
                headers={"Authorization": f"Bearer {auth_token}"},
            )
    except (httpx.TimeoutException, httpx.ConnectError) as exc:
        logger.warning("Media artifact check timed out for phase %s: %s", phase, exc)
        return {"exists": False, "modified_at": None, "age_days": None}

    if resp.status_code == 404:
        return {"exists": False, "modified_at": None, "age_days": None}

    if not resp.is_success:
        logger.warning("Media check returned %d for phase %s", resp.status_code, phase)
        return {"exists": False, "modified_at": None, "age_days": None}

    data = resp.json()
    raw = data.get("__Updated") or data.get("Updated") or data.get("updated")

    modified_at: str | None = None
    age_days: int | None = None
    if raw:
        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            modified_at = dt.isoformat()
            age_days = (datetime.now(timezone.utc) - dt).days
        except ValueError:
            pass

    return {"exists": True, "modified_at": modified_at, "age_days": age_days}


# ── Word document generation ──────────────────────────────────────────────────

def generate_phase_docx(
    phase: str,
    title: str,
    tenant: str,
    site: str,
    sections: list[dict],
) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)

    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    meta = doc.add_paragraph()
    meta.add_run(
        f"Phase: {phase}  |  Site: {tenant} / {site}  |  Generated: {now_str}"
    ).italic = True
    doc.add_paragraph()

    for section in sections:
        heading = section.get("heading", "")
        content = section.get("content", "")
        subsections = section.get("subsections", [])

        if heading:
            doc.add_heading(heading, level=2)
        for line in (content or "").split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        for sub in subsections:
            if sub.get("heading"):
                doc.add_heading(sub["heading"], level=3)
            for line in (sub.get("content") or "").split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Media folder management ───────────────────────────────────────────────────

async def _ssc_get_item_id(cm_host: str, path: str, auth_token: str) -> str | None:
    """Return the bare GUID (no braces) for the item at *path*, or None."""
    try:
        async with httpx.AsyncClient(timeout=8) as http:
            resp = await http.get(
                f"{cm_host}/sitecore/api/ssc/item",
                params={"database": "master", "path": path},
                headers={"Authorization": f"Bearer {auth_token}"},
            )
        if resp.status_code == 200:
            raw_id = resp.json().get("ItemID", "")
            return raw_id.strip("{}") if raw_id else None
        return None
    except Exception:
        return None


async def _ssc_item_exists(cm_host: str, path: str, auth_token: str) -> bool:
    """Return True if a Sitecore item exists at *path* (SSC API)."""
    return (await _ssc_get_item_id(cm_host, path, auth_token)) is not None


async def _create_media_folder(
    cm_host: str, parent_path: str, folder_name: str, auth_token: str
) -> bool:
    """Create a media-library folder via the GraphQL authoring API."""
    # CreateItemInput.parent is typed ID! — pass the bare GUID obtained from SSC.
    # Declaring the variable as String! (even with a valid GUID string) is rejected.
    # templateId is inlined as a literal to avoid Guid scalar parseValue issues.
    parent_id = await _ssc_get_item_id(cm_host, parent_path, auth_token)
    if not parent_id:
        logger.warning("Cannot create %s/%s: parent not found via SSC", parent_path, folder_name)
        return False

    mutation = f"""
mutation CreateMediaFolder($name: String!, $parent: ID!) {{
  createItem(input: {{
    name: $name
    templateId: "{FOLDER_TEMPLATE_ID}"
    parent: $parent
    language: "en"
  }}) {{
    item {{ itemId path }}
  }}
}}
"""
    variables = {"name": folder_name, "parent": parent_id}
    try:
        async with httpx.AsyncClient(timeout=15) as http:
            resp = await http.post(
                f"{cm_host}/sitecore/api/authoring/graphql/v1",
                json={"query": mutation, "variables": variables},
                headers={
                    "Authorization": f"Bearer {auth_token}",
                    "Content-Type": "application/json",
                },
            )
        if resp.is_success:
            data = resp.json()
            if not data.get("errors"):
                item = (data.get("data") or {}).get("createItem", {}).get("item")
                logger.info("Created media folder: %s", (item or {}).get("path", folder_name))
                return True
            logger.warning("GraphQL createItem errors for %s/%s: %s", parent_path, folder_name, data["errors"])
        else:
            logger.warning(
                "GraphQL createItem HTTP %d for %s/%s: %s",
                resp.status_code, parent_path, folder_name, resp.text[:300],
            )
    except Exception as exc:
        logger.warning("GraphQL createItem exception for %s/%s: %s", parent_path, folder_name, exc)
    return False


async def ensure_phase_upload_folders(
    tenant: str, site: str, phase: str, auth_token: str
) -> None:
    """Ensure the Content Strategy folder exists in the media library.

    All phase artifacts are stored flat under Content Strategy — no per-phase
    subfolders. The site folder is assumed to exist. Failures are logged but
    never raise — callers should still attempt the upload.
    """
    cm_host = os.environ.get("SITECORE_CM_HOST", "").rstrip("/")
    if not cm_host:
        return

    if phase not in PHASE_ARTIFACT_MAP:
        return

    site_path = f"/sitecore/Media Library/Project/{tenant}/{site}"
    cs_path = f"{site_path}/{CONTENT_STRATEGY_FOLDER}"

    if not await _ssc_item_exists(cm_host, cs_path, auth_token):
        logger.info("Creating '%s' folder under %s", CONTENT_STRATEGY_FOLDER, site_path)
        await _create_media_folder(cm_host, site_path, CONTENT_STRATEGY_FOLDER, auth_token)


# ── Media library upload ──────────────────────────────────────────────────────

def _media_item_path(full_path: str) -> str:
    """Return the itemPath for uploadMedia: media-library-relative, no extension.

    /sitecore/Media Library/Project/t/s/Research/research-brief.docx
      → Project/t/s/Research/research-brief

    uploadMedia rejects the sitecore/media library prefix, leading slashes, and
    file extensions in item names (dots are invalid per this instance's
    InvalidItemNameChars setting). Pass includeExtensionInItemName=true in the
    mutation so Sitecore appends the extension from the uploaded file, keeping
    the resulting item path consistent with build_artifact_media_path.
    """
    path = full_path.lstrip("/")
    prefix = "sitecore/media library/"
    if path.lower().startswith(prefix):
        path = path[len(prefix):]
    # Strip extension from the item name (last path segment).
    if "/" in path:
        folder, _, name = path.rpartition("/")
        if "." in name:
            name = name.rsplit(".", 1)[0]
        path = f"{folder}/{name}"
    elif "." in path:
        path = path.rsplit(".", 1)[0]
    return path


async def upload_artifact_to_media_library(
    tenant: str,
    site: str,
    phase: str,
    docx_bytes: bytes,
    auth_token: str,
) -> dict:
    info = PHASE_ARTIFACT_MAP.get(phase)
    if not info:
        return {"success": False, "error": f"Unknown phase: {phase!r}", "media_path": None, "overwrite": False}

    media_path = build_artifact_media_path(tenant, site, phase)
    cm_host = os.environ.get("SITECORE_CM_HOST", "").rstrip("/")
    if not cm_host:
        return {"success": False, "error": "SITECORE_CM_HOST not configured", "media_path": media_path, "overwrite": False}

    graphql_url = f"{cm_host}/sitecore/api/authoring/graphql/v1"

    overwrite = False
    try:
        check = await check_media_artifact_exists(tenant, site, phase, auth_token)
        overwrite = check["exists"]
    except Exception:
        pass

    try:
        await ensure_phase_upload_folders(tenant, site, phase, auth_token)
    except Exception as exc:
        logger.warning("Folder pre-creation failed (upload may still succeed): %s", exc)

    # Step 1 — obtain a presigned upload URL from the GraphQL authoring API.
    # Do NOT set includeExtensionInItemName: this instance's InvalidItemNameChars
    # includes '.', so appending ".docx" to the item name produces "research-brief docx"
    # (dot → space). The item is created under the bare name (e.g. "research-brief"),
    # which is what build_artifact_media_path and check_media_artifact_exists use.
    mutation = """
mutation UploadMedia($itemPath: String!) {
  uploadMedia(input: {
    itemPath: $itemPath
    overwriteExisting: true
  }) {
    presignedUploadUrl
  }
}
"""
    logger.info("Requesting presigned upload URL for %s", media_path)
    try:
        async with httpx.AsyncClient(timeout=15) as http:
            resp = await http.post(
                graphql_url,
                json={"query": mutation, "variables": {"itemPath": _media_item_path(media_path)}},
                headers={"Authorization": f"Bearer {auth_token}", "Content-Type": "application/json"},
            )
        if not resp.is_success:
            logger.error("uploadMedia mutation HTTP %d: %s", resp.status_code, resp.text[:500])
            return {
                "success": False,
                "error": f"uploadMedia mutation failed: HTTP {resp.status_code} — {resp.text[:500]}",
                "media_path": media_path,
                "overwrite": overwrite,
            }
        body = resp.json()
        if body.get("errors"):
            logger.error("uploadMedia GraphQL errors: %s", body["errors"])
            return {
                "success": False,
                "error": f"uploadMedia mutation error: {body['errors']}",
                "media_path": media_path,
                "overwrite": overwrite,
            }
        presigned_url = body["data"]["uploadMedia"]["presignedUploadUrl"]
    except httpx.TimeoutException:
        return {"success": False, "error": "Timed out requesting presigned upload URL", "media_path": media_path, "overwrite": overwrite}
    except Exception as exc:
        return {"success": False, "error": f"Failed to get presigned upload URL: {exc}", "media_path": media_path, "overwrite": overwrite}

    # Step 2 — upload file bytes to the presigned URL.
    # This Sitecore instance routes presigned URLs through its own CM host, so
    # the Bearer token is still required (unlike true pre-signed S3 URLs).
    logger.info("Uploading %s via presigned URL (overwrite=%s)", info["filename"], overwrite)
    try:
        async with httpx.AsyncClient(timeout=60) as http:
            resp = await http.post(
                presigned_url,
                files={
                    "file": (
                        info["filename"],
                        docx_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers={"Authorization": f"Bearer {auth_token}"},
            )
        if not resp.is_success:
            logger.error("Presigned upload HTTP %d: %s", resp.status_code, resp.text[:500])
            return {
                "success": False,
                "error": f"Presigned upload failed: HTTP {resp.status_code} — {resp.text[:500]}",
                "media_path": media_path,
                "overwrite": overwrite,
            }
    except httpx.TimeoutException:
        return {"success": False, "error": "Media upload timed out after 60 seconds", "media_path": media_path, "overwrite": overwrite}
    except Exception as exc:
        return {"success": False, "error": f"Media upload failed: {exc}", "media_path": media_path, "overwrite": overwrite}

    logger.info("Uploaded %s to %s (overwrite=%s)", info["filename"], media_path, overwrite)
    return {"success": True, "media_path": media_path, "overwrite": overwrite, "error": None}


# ── Artifact text extraction ──────────────────────────────────────────────────

async def download_and_extract_artifact(
    media_path: str, auth_token: str
) -> tuple[bool, str | None, str | None]:
    cm_host = os.environ.get("SITECORE_CM_HOST", "").rstrip("/")
    if not cm_host:
        return False, None, "SITECORE_CM_HOST not configured"

    # Convert media library path to HTTP download URL
    # e.g. /sitecore/media library/Project/... → /-/media/Project/...
    relative = media_path.replace("/sitecore/media library", "").lstrip("/")
    url = f"{cm_host}/-/media/{relative}"

    try:
        async with httpx.AsyncClient(timeout=15) as http:
            resp = await http.get(url, headers={"Authorization": f"Bearer {auth_token}"})
        if resp.status_code == 404:
            return False, None, "Artifact not found at media library path"
        if not resp.is_success:
            return False, None, f"Failed to download artifact: HTTP {resp.status_code}"
        docx_bytes = resp.content
    except httpx.TimeoutException:
        return False, None, "Download timed out"
    except Exception as exc:
        return False, None, f"Download failed: {exc}"

    try:
        doc = Document(io.BytesIO(docx_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return True, text, None
    except Exception as exc:
        return False, None, f"Failed to extract text from document: {exc}"

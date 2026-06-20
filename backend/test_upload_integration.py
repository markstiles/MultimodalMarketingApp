"""
Standalone integration test for the Sitecore Agent API upload endpoint.

Run from the backend directory with your .env loaded:

    cd backend
    python test_upload_integration.py

Iterates through several encoding variations so we can see which one the API accepts.
"""
import asyncio
import io
import json
import os
import sys
from pathlib import Path

# Load .env from the project root
env_path = Path(__file__).parent.parent / ".env"
if env_path.exists():
    for line in env_path.read_text().splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        os.environ.setdefault(key.strip(), value.strip())

import httpx
from docx import Document


# ── helpers ──────────────────────────────────────────────────────────────────

async def get_token() -> str:
    client_id = os.environ.get("SITECORE_CLIENT_ID_AUTOMATION") or os.environ.get("AUTHOR_APP_ID", "")
    client_secret = os.environ.get("SITECORE_CLIENT_SECRET_AUTOMATION") or os.environ.get("AUTHOR_APP_CLIENT_CREDENTIALS", "")
    if not client_id or not client_secret:
        sys.exit("ERROR: Sitecore credentials not found in .env")

    async with httpx.AsyncClient(timeout=10) as http:
        resp = await http.post(
            "https://auth.sitecorecloud.io/oauth/token",
            data={
                "grant_type": "client_credentials",
                "client_id": client_id,
                "client_secret": client_secret,
                "audience": "https://api.sitecorecloud.io",
            },
        )
        resp.raise_for_status()
        return resp.json()["access_token"]


def make_docx() -> bytes:
    doc = Document()
    doc.add_heading("Test Research Brief", level=1)
    doc.add_paragraph("This is a test document for the upload integration test.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def try_upload(label: str, token: str, docx: bytes, files: dict, data: dict = None) -> None:
    url = (
        os.environ.get("SITECORE_AGENTS_API_BASE_URL", "https://edge-platform.sitecorecloud.io/stream/ai-agent-api").rstrip("/")
        + "/api/v1/assets/upload"
    )
    print(f"\n{'='*60}")
    print(f"VARIANT: {label}")
    print(f"URL: {url}")

    kwargs = dict(
        headers={"Authorization": f"Bearer {token}"},
        files=files,
    )
    if data:
        kwargs["data"] = data

    # Show what we're sending
    req = httpx.Request("POST", url, **kwargs)
    req.read()
    body_preview = req.content.decode("latin-1")
    print("--- MULTIPART BODY (first 600 chars) ---")
    sys.stdout.buffer.write((body_preview[:600] + "\n").encode("utf-8", errors="replace"))

    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(url, **kwargs)

    print(f"--- RESPONSE: HTTP {resp.status_code} ---")
    print(resp.text[:1000])


# ── main ─────────────────────────────────────────────────────────────────────

async def main():
    print("Getting auth token...")
    token = await get_token()
    print(f"Token: {token[:20]}...")

    docx = make_docx()

    url = (
        os.environ.get("SITECORE_AGENTS_API_BASE_URL", "https://edge-platform.sitecorecloud.io/stream/ai-agent-api").rstrip("/")
        + "/api/v1/assets/upload"
    )

    site_name = os.environ.get("SITECORE_PUBLIC_DEFAULT_SITE_NAME", "solovay-s")
    upload_request_dict = {
        "name": "research-brief",
        "itemPath": f"/sitecore/Media Library/Project/velir-studio/{site_name}/Content Strategy/Research",
        "language": "en",
        "extension": "docx",
        "siteName": site_name,
    }
    upload_request_str = json.dumps(upload_request_dict)
    print(f"\nupload_request: {upload_request_str}")

    docx_ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    # Variant 1: files= only, upload_request as (None, str) — current code
    await try_upload(
        "1: files only, (None, str)",
        token, docx,
        files={
            "file": ("research-brief.docx", docx, docx_ct),
            "upload_request": (None, upload_request_str),
        },
    )

    # Variant 2: files= + data=
    await try_upload(
        "2: files + data dict",
        token, docx,
        files={"file": ("research-brief.docx", docx, docx_ct)},
        data={"upload_request": upload_request_str},
    )

    # Variant 3: files= only, upload_request as (None, bytes, "text/plain")
    await try_upload(
        "3: files only, (None, bytes, text/plain)",
        token, docx,
        files={
            "file": ("research-brief.docx", docx, docx_ct),
            "upload_request": (None, upload_request_str.encode(), "text/plain"),
        },
    )

    # Variant 4: files= only, upload_request as (None, bytes, "application/json")
    await try_upload(
        "4: files only, (None, bytes, application/json)",
        token, docx,
        files={
            "file": ("research-brief.docx", docx, docx_ct),
            "upload_request": (None, upload_request_str.encode(), "application/json"),
        },
    )

    # Variant 5: upload_request BEFORE file in dict (order preserved in Python 3.7+)
    await try_upload(
        "5: upload_request FIRST, then file",
        token, docx,
        files={
            "upload_request": (None, upload_request_str),
            "file": ("research-brief.docx", docx, docx_ct),
        },
    )

    # Variant 6: name includes extension (as in spec example "homeimage.jpg")
    upload_request_with_ext = json.dumps({**upload_request_dict, "name": "research-brief.docx"})
    await try_upload(
        "6: name WITH extension",
        token, docx,
        files={
            "file": ("research-brief.docx", docx, docx_ct),
            "upload_request": (None, upload_request_with_ext),
        },
    )

    # Variant 7: simplest possible itemPath (just the folder, no project structure)
    upload_request_simple = json.dumps({
        "name": "research-brief",
        "itemPath": "/sitecore/Media Library/Files",
        "language": "en",
        "extension": "docx",
        "siteName": site_name,
    })
    await try_upload(
        "7: simplest itemPath (/sitecore/Media Library/Files)",
        token, docx,
        files={
            "file": ("research-brief.docx", docx, docx_ct),
            "upload_request": (None, upload_request_simple),
        },
    )


    # ── Diagnostic variants ──────────────────────────────────────────────────

    # D1: No auth header → expect 401 (confirms our auth IS working if other calls get 400)
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D1: No Authorization header")
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(
            url,
            files={
                "file": ("research-brief.docx", docx, docx_ct),
                "upload_request": (None, upload_request_str),
            },
        )
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D2: Send only `file`, no `upload_request` → should say upload_request is missing
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D2: Only 'file' field, no 'upload_request'")
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("research-brief.docx", docx, docx_ct)},
        )
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D3: Send only `upload_request`, no `file` → should say file is missing
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D3: Only 'upload_request' field, no 'file'")
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            data={"upload_request": upload_request_str},
        )
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D4: Empty body (no fields at all) → should say both fields are missing
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D4: No body at all")
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(url, headers={"Authorization": f"Bearer {token}"})
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D5: Tiny fake file (b"hello") instead of real DOCX bytes
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D5: Tiny fake file content (b'hello')")
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            files={
                "file": ("research-brief.docx", b"hello", docx_ct),
                "upload_request": (None, upload_request_str),
            },
        )
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D6: Decode and show token claims (no network call)
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D6: JWT token payload (middle segment)")
    import base64
    payload_b64 = token.split(".")[1]
    payload_b64 += "=" * (-len(payload_b64) % 4)
    payload = json.loads(base64.urlsafe_b64decode(payload_b64))
    print(json.dumps({k: v for k, v in payload.items() if k in ("aud", "scope", "sub", "iss", "azp")}, indent=2))

    # D7: upload_request as a QUERY PARAMETER, file in multipart
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D7: upload_request as query param, file in multipart")
    import urllib.parse
    qp_url = url + "?" + urllib.parse.urlencode({"upload_request": upload_request_str})
    async with httpx.AsyncClient(timeout=30) as http:
        resp = await http.post(
            qp_url,
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("research-brief.docx", b"hello", docx_ct)},
        )
    print(f"HTTP {resp.status_code}: {resp.text[:400]}")

    # D8: Fetch OpenAPI spec (no auth needed — public docs endpoint)
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D8: Fetch OpenAPI spec")
    api_base = os.environ.get("SITECORE_AGENTS_API_BASE_URL", "https://edge-platform.sitecorecloud.io/stream/ai-agent-api").rstrip("/")
    for docs_path in ("/openapi.json", "/docs", "/swagger.json", "/redoc"):
        async with httpx.AsyncClient(timeout=10) as http:
            r = await http.get(api_base + docs_path)
        print(f"  {docs_path}: HTTP {r.status_code} ({len(r.content)} bytes)")
        if r.status_code == 200 and docs_path == "/openapi.json":
            spec = r.json()
            upload_endpoint = spec.get("paths", {}).get("/api/v1/assets/upload", {})
            print("  upload endpoint:", json.dumps(upload_endpoint, indent=2)[:1000])

    # D9: Try using the `requests` library (sync, different multipart implementation)
    print(f"\n{'='*60}")
    print("DIAGNOSTIC D9: requests library multipart")
    try:
        import requests as req_lib
        resp9 = req_lib.post(
            url,
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("research-brief.docx", b"hello", docx_ct)},
            data={"upload_request": upload_request_str},
        )
        print(f"HTTP {resp9.status_code}: {resp9.text[:400]}")
    except ImportError:
        print("  requests library not installed, skipping")


if __name__ == "__main__":
    asyncio.run(main())
